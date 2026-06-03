# -*- coding: utf-8 -*-
"""
生成「中美青年创客大赛」参赛作品说明材料 Word 文档。
作品：HoloGrip —— 免穿戴式 AR 教辅平台
内容基于项目真实源码（手势追踪 / Arcball / 表达式引擎 / AI 识图 / Rust 反代）撰写。
采用标准作品说明书章节结构。
"""
import os
from docx import Document
from docx.shared import Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.section import WD_SECTION
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

# ---------------- 主题色 ----------------
PRIMARY = RGBColor(0x1F, 0x3A, 0x5F)   # 深蓝（一级标题）
ACCENT  = RGBColor(0x2E, 0x74, 0xB5)   # 亮蓝（二级标题 / 强调）
DARK    = RGBColor(0x26, 0x2A, 0x2E)   # 正文
GREY    = RGBColor(0x5B, 0x61, 0x68)   # 副文字
LIGHT_FILL = "EAF1F8"
ALT_FILL   = "F5F8FB"
HEAD_FILL  = "1F3A5F"

CN_TITLE = "微软雅黑"
CN_BODY  = "宋体"
EN_FONT  = "Cambria"
MONO     = "Consolas"


def _set_font(run, cn, en, size, bold, color, italic=False):
    run.font.name = en
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.italic = italic
    if color is not None:
        run.font.color.rgb = color
    rPr = run._element.get_or_add_rPr()
    rFonts = rPr.find(qn('w:rFonts'))
    if rFonts is None:
        rFonts = OxmlElement('w:rFonts')
        rPr.append(rFonts)
    rFonts.set(qn('w:ascii'), en)
    rFonts.set(qn('w:hAnsi'), en)
    rFonts.set(qn('w:eastAsia'), cn)


def para(doc, text="", size=12, cn=CN_BODY, en=EN_FONT, bold=False, color=DARK,
         align=None, before=0, after=8, line=1.6, first_indent=True, italic=False):
    p = doc.add_paragraph()
    pf = p.paragraph_format
    pf.space_before = Pt(before)
    pf.space_after = Pt(after)
    pf.line_spacing = line
    if align is not None:
        p.alignment = align
    if first_indent and align is None:
        pf.first_line_indent = Pt(size * 2)
    if text:
        r = p.add_run(text)
        _set_font(r, cn, en, size, bold, color, italic)
    return p


def h1(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(18)
    p.paragraph_format.space_after = Pt(8)
    r = p.add_run(text)
    _set_font(r, CN_TITLE, EN_FONT, 17, True, PRIMARY)
    pPr = p._p.get_or_add_pPr()
    pbdr = OxmlElement('w:pBdr')
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'single')
    bottom.set(qn('w:sz'), '8')
    bottom.set(qn('w:space'), '5')
    bottom.set(qn('w:color'), '2E74B5')
    pbdr.append(bottom)
    pPr.append(pbdr)
    return p


def h2(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.space_after = Pt(5)
    r = p.add_run(text)
    _set_font(r, CN_TITLE, EN_FONT, 13.5, True, PRIMARY)
    return p


def h3(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after = Pt(3)
    r = p.add_run(text)
    _set_font(r, CN_TITLE, EN_FONT, 12, True, ACCENT)
    return p


def bullet(doc, lead, text, size=12):
    p = doc.add_paragraph()
    pf = p.paragraph_format
    pf.space_after = Pt(5)
    pf.line_spacing = 1.5
    pf.left_indent = Pt(size * 2)
    pf.first_line_indent = Pt(0)
    if lead:
        r1 = p.add_run(lead + "：")
        _set_font(r1, CN_TITLE, EN_FONT, size, True, PRIMARY)
        r2 = p.add_run(text)
        _set_font(r2, CN_BODY, EN_FONT, size, False, DARK)
    else:
        r = p.add_run(text)
        _set_font(r, CN_BODY, EN_FONT, size, False, DARK)
    return p


def numbered(doc, text, size=12):
    p = doc.add_paragraph()
    pf = p.paragraph_format
    pf.space_after = Pt(5)
    pf.line_spacing = 1.5
    pf.left_indent = Pt(size * 2)
    pf.first_line_indent = Pt(0)
    r = p.add_run(text)
    _set_font(r, CN_BODY, EN_FONT, size, False, DARK)
    return p


def figure(doc, caption):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after = Pt(3)
    r = p.add_run("［ 此处插入配图：" + caption + " ］")
    _set_font(r, CN_TITLE, EN_FONT, 10.5, False, ACCENT)
    pPr = p._p.get_or_add_pPr()
    pbdr = OxmlElement('w:pBdr')
    for edge in ('top', 'bottom', 'left', 'right'):
        e = OxmlElement('w:' + edge)
        e.set(qn('w:val'), 'dashed')
        e.set(qn('w:sz'), '6')
        e.set(qn('w:space'), '10')
        e.set(qn('w:color'), '9DB8D2')
        pbdr.append(e)
    pPr.append(pbdr)
    return p


def _shade(cell, hexc):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), hexc)
    tcPr.append(shd)


def _vcenter(cell):
    tcPr = cell._tc.get_or_add_tcPr()
    va = OxmlElement('w:vAlign')
    va.set(qn('w:val'), 'center')
    tcPr.append(va)


def cell_text(cell, text, bold=False, color=DARK, size=10.5, cn=CN_BODY,
              align=WD_ALIGN_PARAGRAPH.LEFT, white=False):
    cell.text = ""
    p = cell.paragraphs[0]
    p.alignment = align
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after = Pt(2)
    p.paragraph_format.line_spacing = 1.3
    r = p.add_run(text)
    c = RGBColor(0xFF, 0xFF, 0xFF) if white else color
    _set_font(r, cn, EN_FONT, size, bold, c)
    _vcenter(cell)


def table(doc, headers, rows, widths=None):
    t = doc.add_table(rows=1, cols=len(headers))
    t.style = 'Table Grid'
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    hdr = t.rows[0].cells
    for i, htext in enumerate(headers):
        cell_text(hdr[i], htext, bold=True, white=True, cn=CN_TITLE,
                  align=WD_ALIGN_PARAGRAPH.CENTER, size=10.5)
        _shade(hdr[i], HEAD_FILL)
    for ri, row in enumerate(rows):
        cells = t.add_row().cells
        fill = LIGHT_FILL if ri % 2 == 0 else "FFFFFF"
        for ci, val in enumerate(row):
            al = WD_ALIGN_PARAGRAPH.CENTER if (widths and ci == 0) else WD_ALIGN_PARAGRAPH.LEFT
            cell_text(cells[ci], val, align=al,
                      bold=(ci == 0 and widths is not None))
            _shade(cells[ci], fill)
    if widths:
        for ci, w in enumerate(widths):
            for row in t.rows:
                row.cells[ci].width = Cm(w)
    doc.add_paragraph().paragraph_format.space_after = Pt(2)
    return t


def code_block(doc, lines):
    """等宽字体代码/数据块。"""
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(0.6)
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(8)
    p.paragraph_format.line_spacing = 1.3
    pPr = p._p.get_or_add_pPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear'); shd.set(qn('w:color'), 'auto'); shd.set(qn('w:fill'), 'F2F4F7')
    pPr.append(shd)
    pbdr = OxmlElement('w:pBdr')
    left = OxmlElement('w:left')
    left.set(qn('w:val'), 'single'); left.set(qn('w:sz'), '18'); left.set(qn('w:space'), '6'); left.set(qn('w:color'), '2E74B5')
    pbdr.append(left)
    pPr.append(pbdr)
    for i, line in enumerate(lines):
        r = p.add_run(line)
        _set_font(r, MONO, MONO, 9.5, False, RGBColor(0x33, 0x3A, 0x44))
        if i < len(lines) - 1:
            r.add_break()
    return p


# ============================================================
#  文档构建
# ============================================================
doc = Document()

sec = doc.sections[0]
sec.page_width = Cm(21.0)
sec.page_height = Cm(29.7)
sec.top_margin = Cm(2.5)
sec.bottom_margin = Cm(2.3)
sec.left_margin = Cm(2.7)
sec.right_margin = Cm(2.7)

normal = doc.styles['Normal']
normal.font.name = EN_FONT
normal.font.size = Pt(12)
normal._element.rPr.rFonts.set(qn('w:eastAsia'), CN_BODY)

print("脚本骨架就绪")


# ---------------- 页脚页码 ----------------
def add_page_number(section):
    footer = section.footer
    p = footer.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run()
    fldStart = OxmlElement('w:fldChar'); fldStart.set(qn('w:fldCharType'), 'begin')
    instr = OxmlElement('w:instrText'); instr.set(qn('xml:space'), 'preserve'); instr.text = 'PAGE'
    fldEnd = OxmlElement('w:fldChar'); fldEnd.set(qn('w:fldCharType'), 'end')
    run._r.append(fldStart); run._r.append(instr); run._r.append(fldEnd)
    _set_font(run, CN_BODY, EN_FONT, 9, False, GREY)


# ============================================================
#  封面
# ============================================================
for _ in range(2):
    para(doc, "", after=0, first_indent=False)

para(doc, "中美青年创客大赛", size=15, cn=CN_TITLE, bold=False, color=GREY,
     align=WD_ALIGN_PARAGRAPH.CENTER, after=4)
para(doc, "China-U.S. Young Maker Competition", size=10.5, cn=EN_FONT, color=GREY,
     align=WD_ALIGN_PARAGRAPH.CENTER, after=2)
para(doc, "参 赛 作 品 说 明 书", size=13, cn=CN_TITLE, color=GREY,
     align=WD_ALIGN_PARAGRAPH.CENTER, after=46)

para(doc, "HoloGrip", size=42, cn=CN_TITLE, bold=True, color=PRIMARY,
     align=WD_ALIGN_PARAGRAPH.CENTER, after=6)
para(doc, "免穿戴式 AR 教辅平台", size=23, cn=CN_TITLE, bold=True, color=ACCENT,
     align=WD_ALIGN_PARAGRAPH.CENTER, after=6)
para(doc, "A Wearable-Free Augmented-Reality Teaching-Assistant Platform",
     size=11.5, cn=EN_FONT, color=GREY, align=WD_ALIGN_PARAGRAPH.CENTER, after=30)

para(doc, "无需头显，无需手套，无需触摸屏——一个普通摄像头，", size=13, cn=CN_TITLE,
     color=DARK, align=WD_ALIGN_PARAGRAPH.CENTER, after=2, first_indent=False)
para(doc, "让抽象的数学“看得见、转得动、画得出”。", size=13, cn=CN_TITLE,
     color=DARK, align=WD_ALIGN_PARAGRAPH.CENTER, after=40, first_indent=False)

figure(doc, "作品主界面 / Logo 展示图")
para(doc, "", after=24, first_indent=False)

cover = table(
    doc,
    ["项 目", "内 容"],
    [
        ["作品名称", "HoloGrip —— 免穿戴式 AR 教辅平台"],
        ["作品类别", "教育科技 / 人工智能应用 / 人机交互"],
        ["技术关键词", "计算机视觉、手势识别、增强现实、多模态大模型、三维可视化"],
        ["运行形态", "Web 网页端 + Windows 桌面应用（同源双端）"],
        ["参赛组别", "（待填写）"],
        ["所属赛区", "（待填写）"],
        ["团队名称", "（待填写）"],
        ["团队成员", "（待填写）"],
        ["指导教师", "（待填写）"],
        ["开源地址", "（待填写）"],
        ["完成日期", "2026 年"],
    ],
    widths=[3.8, 11.7],
)

doc.add_page_break()
add_page_number(doc.sections[0])
print("封面就绪")


# ============================================================
#  一、作品概述
# ============================================================
h1(doc, "一、作品概述")

para(doc,
     "HoloGrip 是一款面向中学与高校数学课堂的免穿戴式增强现实（AR）教辅平台。"
     "它的核心理念可以用一句话概括：把昂贵的空间交互体验，压缩进一个普通摄像头里。"
     "用户无需佩戴任何头显、数据手套或专用传感器，也无需触摸屏幕——只要面对一颗普通的电脑摄像头，"
     "在空中用双手捏合、旋转、缩放，就能直接操控悬浮在屏幕中的三维几何体，"
     "在虚拟白板上书写推导，让原本停留在课本插图里的抽象数学“立”起来、“动”起来。")

para(doc,
     "平台围绕数学教学的真实场景，构建了五大功能模块：空间 AR 几何探索、AI 几何识图、"
     "超级电子白板、函数探究器与三维空间计算器。它们既能独立使用，也能在一堂课中无缝衔接——"
     "老师可以先用 AI 识图把课本上的立体几何题一键还原成可旋转的三维模型，"
     "再切到 AR 模式让学生隔空“转动”这个模型观察各个面的位置关系，"
     "最后在电子白板上用虚拟尺规完成作图与推导。")

para(doc,
     "在工程实现上，HoloGrip 基于同一份 React 代码同时产出 Web 网页端与 Windows 桌面应用，"
     "并配套了一套用 Rust 编写的安全反向代理服务，把 AI 大模型密钥牢牢锁在服务端。"
     "整套系统强调“轻量、普惠、可落地”：不依赖专用硬件，一台带摄像头的普通电脑即可运行，"
     "让缺乏 AR/VR 设备预算的学校也能用得起空间交互式数学教学。")

figure(doc, "五大功能模块总览拼图")

# ============================================================
#  二、研发背景与意义
# ============================================================
h1(doc, "二、研发背景与意义")

h2(doc, "2.1  抽象数学的“看不见”之痛")
para(doc,
     "立体几何、函数图像、空间向量是中学与大学数学公认的难点。其根源在于“维度落差”："
     "课本和黑板都是二维平面，而这些知识本质上是三维甚至更高维的空间结构。"
     "学生需要在脑海中完成“从平面图脑补出立体形状”的转换，这种空间想象能力恰恰是初学者最欠缺的。"
     "一张静止的斜二测画法示意图，无法让学生看到“换个角度棱与面会如何遮挡”，"
     "理解便容易卡在“想不出来”这一关。")

h2(doc, "2.2  现有交互手段的局限")
para(doc,
     "为了把数学“可视化”，业界尝试过多条路径，但每条都有明显短板：", after=4)
bullet(doc, "传统实物教具",
       "只能展示固定形状，无法动态生成任意题目里的几何体，也无法与函数、计算联动。")
bullet(doc, "鼠标拖拽的三维软件",
       "交互被束缚在二维鼠标上，旋转操作反直觉，且与“在黑板上讲课”的教学动作割裂。")
bullet(doc, "VR / AR 头显方案",
       "沉浸感强，但设备昂贵、需要逐人佩戴、存在眩晕问题，且不适合一对多的课堂讲授场景，"
       "对经费有限的普通学校几乎不可行。")

h2(doc, "2.3  我们的切入点：免穿戴")
para(doc,
     "HoloGrip 选择了一条被低估的技术路线——用纯软件的计算机视觉，把空间手势交互的门槛降到“一个摄像头”。"
     "这一选择直接回应了教育公平这一中美双方共同关注的议题：当沉浸式交互不再绑定昂贵硬件，"
     "优质的数学可视化教学就有机会从一线城市的实验校，走向预算紧张的县域学校与乡村课堂。"
     "“免穿戴”不只是一个技术参数，更是一种让先进教育工具触手可及的设计立场。")

# ============================================================
#  三、作品创新点
# ============================================================
h1(doc, "三、作品创新点")

h2(doc, "3.1  零硬件依赖的空中手势交互")
para(doc,
     "作品最核心的创新，是仅凭一颗普通摄像头就实现了媲美专业设备的双手空中操控。"
     "系统借助 MediaPipe 手部关键点模型实时捕捉双手的 21 个关节点，"
     "再经过一套自研的手势追踪与平滑层，把抖动的原始识别结果转化为稳定、跟手的交互信号。"
     "用户用拇指与食指“捏合”即可点击，单手平移驱动模型旋转，双手开合控制缩放、双手转腕控制翻滚——"
     "整套交互语言借鉴了 Vision Pro 的隔空操作直觉，却不需要任何一件穿戴设备。")

h2(doc, "3.2  AI 一键还原立体几何题")
para(doc,
     "针对中国数学教材普遍采用的斜二测画法，作品设计了专用的多模态识图链路。"
     "学生只需上传一张立体几何题的截图，系统通过精心设计的提示词约束大模型，"
     "自动测算图中长宽高比例、区分实线与虚线判断前后遮挡、剔除辅助线与动点，"
     "最终输出带顶点标签的结构化三维数据，并在画布上还原成可旋转的立体模型。"
     "把“课本上的死图”变成“手里能转的活模型”，这一步打通了纸面习题与空间交互之间的鸿沟。")

h2(doc, "3.3  一份代码、双端落地的工程范式")
para(doc,
     "作品用同一套 React 代码库，通过运行时平台探测，同时编译出 Web 网页端和 Windows 桌面安装包。"
     "Web 端便于在机房、平板上即开即用；桌面端则提供本地模型库与更流畅的离线体验。"
     "配合 Rust 反向代理的密钥隔离设计，作品在“易部署”与“高安全”之间取得了少见的平衡。")

h2(doc, "3.4  多学科交叉的融合创新")
para(doc,
     "HoloGrip 的诞生本身就是跨界融合的产物：它把计算机视觉与信号处理（卡尔曼滤波、One-Euro 滤波）、"
     "计算机图形学（Arcball 球面旋转）、编译原理（自研表达式引擎）、人机交互设计与数学教育学，"
     "汇聚到同一个产品中。每一个技术选择背后，都对应着一个真实的教学痛点。")

print("第一至三章就绪")


# ============================================================
#  四、功能介绍
# ============================================================
h1(doc, "四、功能介绍")

para(doc,
     "HoloGrip 通过底部 Dock 栏在五大模块间一键切换。下表先给出总览，随后逐一展开。", after=6)

table(
    doc,
    ["模块", "核心能力", "典型场景"],
    [
        ["空间 AR 几何探索", "摄像头手势隔空旋转 / 缩放 / 翻滚三维模型", "讲解棱锥、棱柱各面位置关系"],
        ["AI 几何识图", "上传题图自动还原带标注的三维结构", "把课本立体几何题变为活模型"],
        ["超级电子白板", "虚拟尺规作图 + 画笔 + 在模型表面书写", "课堂板书、几何作图演示"],
        ["函数探究器", "自研表达式引擎 + 滑动条参数联动", "探究函数图像、零点与极值"],
        ["三维空间计算器", "以几何对象可视化呈现运算过程", "直观理解运算的几何意义"],
    ],
    widths=[3.6, 6.4, 5.5],
)

h2(doc, "4.1  空间 AR 几何探索")
para(doc,
     "进入 AR 模式后，摄像头画面成为背景，三维几何模型悬浮其上。用户的双手化作两个空间光标："
     "左手作为“控制手”，捏合即可点击界面按钮、缩放模型；右手作为“书写手 / 连线手”。"
     "单手移动时，系统按 Arcball 算法把光标轨迹映射为模型的自由旋转；"
     "双手同时捏合并开合，模型随之放大缩小；双手连线角度的变化，则驱动模型绕屏幕法线翻滚。"
     "所有动作都在空中完成，真正实现“隔空把玩几何体”。")
figure(doc, "AR 模式下双手操控三维棱锥的实拍画面")

h2(doc, "4.2  AI 几何识图")
para(doc,
     "在识图入口上传一张立体几何题截图，平台会调用多模态大模型进行结构化解析，"
     "先输出一段“测量与剔除”的推理过程，再给出几何体名称、带标签的顶点坐标、面与棱的拓扑关系。"
     "前端随后将任意比例的顶点集合归一化到统一视觉尺度，并把多边形面三角化，交由三维引擎渲染。"
     "几秒之内，一道纸面习题就变成了可在 AR 空间中旋转观察的立体模型。")
figure(doc, "AI 识图：题目截图 → 还原的三维模型对比图")

h2(doc, "4.3  超级电子白板")
para(doc,
     "白板模块内置可吸附的虚拟直尺、三角板、量角器和圆规，配合画笔与橡皮擦，"
     "完整复刻了老师在黑板上作图的工具链。沿尺子边缘可印出直线，沿量角器、圆规可印出角度与圆弧。"
     "更进一步，笔迹还能直接书写在三维模型表面，并存储于模型的局部坐标系——"
     "模型旋转时，写下的标注会跟着一起转，仿佛真的“刻”在了立体表面上。")
figure(doc, "电子白板：虚拟尺规作图与模型表面书写")

h2(doc, "4.4  函数探究器")
para(doc,
     "函数模块由一套完全自研的表达式引擎驱动，支持四则运算、乘方、隐式乘法（如 2x、3sin(x)）、"
     "复合函数以及二十余种常用数学函数。用户输入函数表达式后，系统不仅实时绘制图像，"
     "还能用数值方法自动求解零点与极值，并通过滑动条调节参数、观察图像的动态变化，"
     "把“参数如何影响函数形态”这一抽象关系变得可视、可调、可感。")
figure(doc, "函数探究器：参数滑动条与实时图像、零点 / 极值标注")

h2(doc, "4.5  三维空间计算器")
para(doc,
     "区别于普通计算器只给出数字结果，三维空间计算器把运算过程以可视化的几何对象呈现，"
     "帮助学生理解数与形之间的联系，让计算不再是黑箱里的按键，而是看得见的几何变换。")
figure(doc, "三维空间计算器界面")

print("第四章就绪")


# ============================================================
#  五、技术方案与实现
# ============================================================
h1(doc, "五、技术方案与实现")

h2(doc, "5.1  总体架构")
para(doc,
     "系统整体遵循“前端富客户端 + 安全反代隔离 + 上游 AI 网关”的分层结构。"
     "Web 端与桌面端共享同一份前端产物，区别仅在运行容器与对外通信路径：", after=4)
bullet(doc, "用户交互层",
       "门户首页（原生 JS/CSS）启动主程序；主程序由 React 19 + React-Three-Fiber + "
       "Tailwind + Zustand 构成，承载 AR、白板、函数、计算器等模块。")
bullet(doc, "运行容器层",
       "Web 端经 Nginx 提供静态资源并反代 API；桌面端为 Tauri 容器（WebView2 + Rust），"
       "带自绘标题栏与本地模型库，编译产物为约 12 MB 的单一安装包，无需额外安装运行时。")
bullet(doc, "安全反代层",
       "独立的 Rust（axum）反向代理服务，负责鉴权、限流、密钥注入，是前端与上游 AI 之间的唯一通道。")
bullet(doc, "上游 AI 网关",
       "多模态大模型服务，承担立体几何图片的结构化理解。")

table(
    doc,
    ["层次", "关键技术选型", "职责"],
    [
        ["前端框架", "React 19 + TypeScript", "组件化 UI 与状态驱动渲染"],
        ["三维渲染", "Three.js + React-Three-Fiber", "几何体建模、光照、材质与交互"],
        ["手势识别", "MediaPipe Tasks-Vision (GPU)", "实时提取双手 21 个关键点"],
        ["状态管理", "Zustand", "手势 / 模型 / UI / 笔迹 / 主题单一真相源"],
        ["桌面封装", "Tauri 2.x + Rust", "打包 Windows 桌面应用"],
        ["后端反代", "Rust + axum + tower", "鉴权、限流、密钥隔离、流式转发"],
    ],
    widths=[3.2, 6.3, 6.0],
)
figure(doc, "系统总体分层架构示意图")

h2(doc, "5.2  手势追踪：让识别结果“稳得住、跟得上”")
para(doc,
     "摄像头识别的原始坐标天然充满抖动，且常有短时丢失、多人入镜、左右手混淆等问题。"
     "若直接使用，模型会不停颤抖、光标会乱跳，根本无法用于教学。"
     "为此，作品在 MediaPipe 之上自研了一整套手势追踪层，核心由三项技术构成：")

h3(doc, "（1）常加速度卡尔曼滤波")
para(doc,
     "系统为每只手在 X、Y 轴上各维护一个“位置–速度–加速度”三维状态，对每帧观测做"
     "“预测—更新”闭环，输出方差约束下的最优估计。更关键的是采用了自适应噪声策略："
     "手快速挥动时放大过程噪声以紧紧跟手，手缓慢移动时收紧噪声以滤除微抖，"
     "兼顾了灵敏度与平滑度。")

h3(doc, "（2）Coast 滑行与施密特触发器")
para(doc,
     "当识别短时丢失（默认 250 毫秒内），系统不会让光标瞬间消失或跳变，"
     "而是依据上一帧的速度与加速度做物理外推，平滑“滑行”过这段空窗；"
     "一旦外推距离超出合理阈值则主动放弃，避免乱飞。"
     "捏合判定则采用施密特触发器（双阈值迟滞）：未捏合需拉近到 0.045 才触发，"
     "已捏合需拉远到 0.06 才释放，彻底消除了临界点的反复抖动。")

h3(doc, "（3）主用户锁定与左右手防跳")
para(doc,
     "面向一对多的课堂，系统以最先入镜的双手为“主用户”，旁观者的手会被忽略，"
     "锁定超时后才允许新用户接管，避免镜头里多人时控制权乱跳。"
     "左右手的判定则结合 handedness 多帧投票与位置代价匹配双重保险，并对身份做槽位锁定，"
     "从根本上解决了上一版“双手靠近屏幕中线时光标互换”的问题。")
para(doc,
     "下面是手势追踪层中可调参数的真实节选，体现了系统对交互手感的精细打磨：", before=4, after=2)
code_block(doc, [
    "coastDurationMs:        250    // 丢失多少毫秒内仍做外推滑行",
    "maxCoastExtrapolation:  0.35   // coast 期最大外推距离 (NDC)",
    "handednessVoteWindow:   10     // 左右手投票窗口帧数",
    "primaryLockTimeoutMs:   1500   // 主用户锁定持久度",
    "捏合迟滞阈值:           0.045 / 0.06   // 施密特触发器双阈值",
    "微动死区:               速度<0.12 且位移<0.0018 → 强制锁定静止",
])
figure(doc, "手势数据流水线：原始识别 → 滤波平滑 → 稳定光标")

h2(doc, "5.3  三维旋转：路径无关的 Arcball 算法")
para(doc,
     "模型旋转采用图形学中的经典 Arcball（虚拟轨迹球）方案。系统把屏幕光标投影到一个单位球面上，"
     "以前后两帧投影点的叉积为旋转轴、点积的反余弦为旋转角，生成增量四元数。"
     "这种方法的优势是“路径无关”——无论手怎么绕，绕一圈回到原点，模型姿态也会精确复位，"
     "不会产生累积漂移。投影使用 Bell's Trackball 公式（Holroyd 改进版），保证球面边缘平滑过渡。")
para(doc,
     "在输入端，系统还叠加了 One-Euro 滤波器：慢速移动时强滤波吃掉抖动，快速移动时弱滤波保留跟手感，"
     "几乎零延迟。再配合单帧最大角度限制与位移死区，让旋转既灵敏又稳定。")

h2(doc, "5.4  函数引擎：自研的表达式编译器")
para(doc,
     "函数探究器没有使用任何第三方公式库，而是自研了一个基于调度场算法（Shunting-Yard）的"
     "表达式编译引擎。它先对输入做词法分析，自动插入隐式乘法（如把 2x 理解为 2*x），"
     "再转换为逆波兰式并编译成快速求值闭包；编译期即对括号匹配、参数个数等做校验。"
     "在此基础上，系统用采样变号加二分细化的数值方法求解函数零点，"
     "用中心差分法求导并定位极大、极小值点。整套引擎轻量、安全（不执行任意代码）、且响应迅速。")
code_block(doc, [
    "输入:  3sin(x) + 2x^2 - 1",
    "  ↓ 词法分析 + 隐式乘法插入",
    "  ↓ Shunting-Yard → 逆波兰式 (RPN)",
    "  ↓ 编译为求值闭包 + 编译期校验",
    "输出:  实时图像 + 零点 + 极值 (滑动条参数联动)",
])

h2(doc, "5.5  安全设计：密钥永不出仓")
para(doc,
     "AI 能力依赖大模型密钥，一旦泄露后果严重。作品的解决方案是：Web 端的前端产物里"
     "绝不包含任何真实密钥。所有 AI 请求先经浏览器向 Rust 反代申请一个短期令牌（token），"
     "再带着令牌请求；真实密钥仅由服务端在转发时注入。该令牌的安全设计包括：", after=4)
bullet(doc, "HMAC-SHA256 签名",
       "令牌由服务端密钥签名，校验采用常量时间比较以防计时攻击。")
bullet(doc, "三重约束",
       "令牌默认 1 小时过期、绑定签发时的客户端 IP、设有调用次数上限，泄露代价远小于密钥本身。")
bullet(doc, "限流与防护",
       "基于 IP 的令牌桶限流、请求体大小限制、上游超时控制、服务重启即令牌全部作废，多重兜底防滥用。")
para(doc,
     "桌面端则由 Tauri 的 Rust 进程持有密钥，前端 bundle 同样搜不到真实密钥。"
     "这套设计让作品既能公开部署演示，又不必担心核心凭据外泄。")
figure(doc, "Web / 桌面双端密钥隔离链路图")

print("第五章就绪")


# ============================================================
#  六、特色与优势
# ============================================================
h1(doc, "六、特色与优势")

h2(doc, "6.1  与同类方案的对比")
table(
    doc,
    ["对比维度", "VR/AR 头显方案", "鼠标三维软件", "HoloGrip"],
    [
        ["硬件成本", "高（专用设备）", "低", "极低（仅需摄像头）"],
        ["空间交互", "强", "弱（受限于鼠标）", "强（双手隔空操作）"],
        ["课堂适配", "需逐人佩戴", "讲授割裂", "一对多、贴合板书"],
        ["上手门槛", "需培训", "操作反直觉", "捏合即点、直觉自然"],
        ["AI 题目还原", "一般不具备", "一般不具备", "支持一键还原"],
        ["部署方式", "受设备限制", "需安装", "网页即开 / 桌面安装"],
    ],
    widths=[3.2, 3.9, 3.4, 4.0],
)

h2(doc, "6.2  核心优势提炼")
bullet(doc, "普惠可及",
       "零专用硬件依赖，把空间交互式教学的成本降到“一台普通电脑”，天然契合教育公平。")
bullet(doc, "稳定可用",
       "三重滤波与追踪算法保证手势识别在真实课堂光照、抖动下依然稳定跟手，不是实验室样品。")
bullet(doc, "场景闭环",
       "识图、AR、白板、函数、计算五模块衔接，覆盖“出题—演示—作图—探究”的完整教学链路。")
bullet(doc, "工程扎实",
       "双端同源、密钥隔离、自研引擎，体现了从交互算法到后端安全的全栈完成度。")

# ============================================================
#  七、应用场景与教育价值
# ============================================================
h1(doc, "七、应用场景与教育价值")

h2(doc, "7.1  典型应用场景")
bullet(doc, "课堂演示",
       "教师在讲台上隔空旋转几何体、调节函数参数，全班同步观看，替代静态投影与实物教具。")
bullet(doc, "学生自学",
       "学生上传习题截图自动还原模型，反复旋转观察，攻克空间想象难关。")
bullet(doc, "作图教学",
       "用虚拟尺规在白板上规范作图，演示尺规作图的每一步，并可在模型表面标注。")
bullet(doc, "探究式学习",
       "通过滑动条实时观察参数对函数图像的影响，把结论性知识变为可亲手探索的过程。")

h2(doc, "7.2  教育价值与社会意义")
para(doc,
     "HoloGrip 直接服务于联合国可持续发展目标中的“优质教育”这一中美共同关注的议题。"
     "它把原本需要昂贵设备才能获得的沉浸式交互体验，下沉为普通摄像头即可运行的轻量工具，"
     "让数学可视化教学不再是少数资源充足学校的专利。对空间想象能力薄弱的学生，"
     "它提供了“看得见、转得动”的直观支架；对缺乏教具与设备预算的学校，"
     "它提供了一条低成本的数字化教学升级路径。这种“用算法替代硬件”的思路，"
     "正是创客精神在教育普惠方向上的一次具体实践。")
figure(doc, "课堂应用场景实拍图")

# ============================================================
#  八、未来规划
# ============================================================
h1(doc, "八、未来规划")
para(doc, "在现有原型基础上，作品后续将沿以下方向迭代完善：", after=4)
numbered(doc, "1. 学科拓展：从立体几何延伸到解析几何、空间向量、概率统计等更多数学主题，并探索物理、化学的三维可视化。")
numbered(doc, "2. 手势丰富化：增加更多语义手势（如抓取、剖切、展开多面体），让交互表达更细腻。")
numbered(doc, "3. 多人协作：支持师生、生生在同一空间中协同操作同一模型，强化课堂互动。")
numbered(doc, "4. 跨平台覆盖：在现有 Web 与 Windows 桌面端基础上，适配平板与更多操作系统。")
numbered(doc, "5. 开源共建：完善技术文档与开发指南，以开源方式开放代码，吸纳社区意见持续打磨。")

# ============================================================
#  九、结语
# ============================================================
h1(doc, "九、结语")
para(doc,
     "HoloGrip 用一颗普通摄像头，重新定义了数学课堂里“人与知识”的交互方式。"
     "它证明了：先进的空间交互体验不一定要绑定昂贵硬件，普惠与创新可以并行不悖。"
     "我们希望以这件作品为起点，让更多学生能够伸出双手，"
     "真正“抓住”那些曾经只存在于想象中的数学之美。")

para(doc, "", after=10, first_indent=False)
endp = doc.add_paragraph()
endp.alignment = WD_ALIGN_PARAGRAPH.CENTER
er = endp.add_run("—— HoloGrip 项目组 ——")
_set_font(er, CN_TITLE, EN_FONT, 11, False, GREY)

# ============================================================
#  保存
# ============================================================
out_path = os.path.join(
    os.path.dirname(os.path.dirname(os.path.abspath(__file__))),
    "HoloGrip-中美青年创客大赛作品说明书.docx",
)
doc.save(out_path)
print("已保存:", out_path)
