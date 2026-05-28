# -*- coding: utf-8 -*-
"""
HoloMath 技术架构文档生成脚本

读取本仓库的源码与配置，输出一份完整的 Word 技术架构说明书。
执行方式:  python scripts/build_arch_doc.py
"""

from __future__ import annotations

import os
from pathlib import Path

from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

ROOT = Path(__file__).resolve().parent.parent
OUT = ROOT / "HoloMath-技术架构文档.docx"

# ────────────────────────────── 样式工具 ──────────────────────────────


def set_cn_font(run, font_name: str = "微软雅黑", size_pt: float = 10.5,
                bold: bool = False, color: tuple[int, int, int] | None = None):
    run.font.name = font_name
    r = run._element
    rPr = r.get_or_add_rPr()
    rFonts = rPr.find(qn("w:rFonts"))
    if rFonts is None:
        rFonts = OxmlElement("w:rFonts")
        rPr.append(rFonts)
    rFonts.set(qn("w:eastAsia"), font_name)
    rFonts.set(qn("w:ascii"), font_name)
    rFonts.set(qn("w:hAnsi"), font_name)
    run.font.size = Pt(size_pt)
    run.bold = bold
    if color:
        run.font.color.rgb = RGBColor(*color)


def add_heading(doc: Document, text: str, level: int = 1):
    """带正确中文字体的标题。"""
    sizes = {0: 22, 1: 18, 2: 14, 3: 12}
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = p.add_run(text)
    set_cn_font(run, font_name="微软雅黑", size_pt=sizes.get(level, 12),
                bold=True, color=(31, 73, 125))
    p.paragraph_format.space_before = Pt(12 if level <= 1 else 6)
    p.paragraph_format.space_after = Pt(6)
    return p


def add_para(doc: Document, text: str, size_pt: float = 10.5,
             bold: bool = False, indent_first: bool = True):
    p = doc.add_paragraph()
    if indent_first:
        p.paragraph_format.first_line_indent = Cm(0.74)
    p.paragraph_format.line_spacing = 1.4
    run = p.add_run(text)
    set_cn_font(run, "宋体", size_pt, bold=bold)
    return p


def add_bullets(doc: Document, items: list[str], size_pt: float = 10.5):
    for it in items:
        p = doc.add_paragraph(style="List Bullet")
        p.paragraph_format.line_spacing = 1.35
        run = p.add_run(it)
        set_cn_font(run, "宋体", size_pt)


def add_code(doc: Document, code: str):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(0.5)
    p.paragraph_format.line_spacing = 1.15
    run = p.add_run(code)
    set_cn_font(run, "Consolas", 9.5, color=(0x33, 0x33, 0x33))
    # 浅灰底纹
    pPr = p._p.get_or_add_pPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), "F4F6FA")
    pPr.append(shd)


def add_table(doc: Document, header: list[str], rows: list[list[str]]):
    table = doc.add_table(rows=1 + len(rows), cols=len(header))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Light Grid Accent 1"
    hdr = table.rows[0].cells
    for i, h in enumerate(header):
        cell = hdr[i]
        cell.text = ""
        run = cell.paragraphs[0].add_run(h)
        set_cn_font(run, "微软雅黑", 10.5, bold=True, color=(255, 255, 255))
        # 头部蓝底
        tcPr = cell._tc.get_or_add_tcPr()
        shd = OxmlElement("w:shd")
        shd.set(qn("w:val"), "clear")
        shd.set(qn("w:color"), "auto")
        shd.set(qn("w:fill"), "1F497D")
        tcPr.append(shd)
    for ri, row in enumerate(rows, start=1):
        for ci, val in enumerate(row):
            cell = table.rows[ri].cells[ci]
            cell.text = ""
            run = cell.paragraphs[0].add_run(val)
            set_cn_font(run, "宋体", 10)
    doc.add_paragraph()  # 表后空行


# ────────────────────────────── 主体内容 ──────────────────────────────


def build():
    doc = Document()

    # 全局页边距
    section = doc.sections[0]
    section.left_margin = Cm(2.4)
    section.right_margin = Cm(2.4)
    section.top_margin = Cm(2.2)
    section.bottom_margin = Cm(2.2)

    # 默认正文样式中文支持
    style = doc.styles["Normal"]
    style.font.name = "宋体"
    style.element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
    style.font.size = Pt(10.5)

    # ============ 封面 ============
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run("HoloMath 技术架构文档")
    set_cn_font(run, "微软雅黑", 28, bold=True, color=(31, 73, 125))
    title.paragraph_format.space_before = Pt(80)

    sub = doc.add_paragraph()
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = sub.add_run("沉浸式空间数学探索工作台")
    set_cn_font(run, "微软雅黑", 16, color=(80, 80, 80))

    sub2 = doc.add_paragraph()
    sub2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = sub2.add_run("基于 React 19 + Three.js + MediaPipe + Tauri + Rust 的多端融合实现")
    set_cn_font(run, "宋体", 11, color=(120, 120, 120))
    sub2.paragraph_format.space_before = Pt(6)

    meta = doc.add_paragraph()
    meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = meta.add_run("\n\n\n版本：v0.1.1\n编制：HoloGrip 项目组\n日期：2026 年")
    set_cn_font(run, "宋体", 11, color=(100, 100, 100))
    meta.paragraph_format.space_before = Pt(150)

    doc.add_page_break()

    # ============ 1. 项目概述 ============
    add_heading(doc, "1  项目概述", 1)

    add_heading(doc, "1.1  产品定位", 2)
    add_para(doc,
        "HoloMath（项目代号 HoloGrip）是一款面向 K12 与高校数学课堂的沉浸式空间几何探索工具。"
        "它通过实时手势识别、空间三维计算器、AI 几何识图与多功能电子白板四个核心模块，"
        "把抽象的数学概念转化为可触可感的三维交互对象，覆盖从课堂演示到自学研讨的全场景。")

    add_heading(doc, "1.2  核心特征", 2)
    add_bullets(doc, [
        "空间 AR 模式：通过摄像头 + MediaPipe HandLandmarker 完成无接触双手手势交互，"
        "支持 Vision Pro 风格的捏合点击、Arcball 旋转、双手缩放与扭转。",
        "AI 几何识图：上传立体几何题截图，调用 Gemini 多模态模型自动还原带顶点标签的"
        "三维结构（顶点 / 面 / 棱），并按斜二测画法精确还原比例。",
        "超级电子白板：内置可吸附的虚拟直尺、三角板、量角器、圆规，支持画笔与橡皮擦，"
        "并可在三维模型表面直接书写笔迹（笔迹存储于模型局部坐标系）。",
        "函数探究器：自研 Shunting-Yard 表达式编译器，支持隐式乘法、复合函数与"
        "极值/零点的数值求解，可与滑动条参数实时联动。",
        "三维空间计算器：以可视化几何对象呈现运算过程的教学型计算器。",
        "双端打包：基于同一份 React 代码同时输出 Web 站点（Nginx 反代部署）与"
        "Tauri 桌面应用（Windows NSIS 安装包，含自绘标题栏与本地模型仓库）。",
    ])

    add_heading(doc, "1.3  功能矩阵概览", 2)
    add_table(doc,
        ["模块", "入口标签", "关键能力", "依赖技术"],
        [
            ["AR 三维空间", "ar_3d", "手势识别、三维模型旋转 / 缩放、表面书写、AI 模型加载",
             "MediaPipe Tasks Vision、@react-three/fiber、Zustand"],
            ["超级白板", "whiteboard", "自由书写、几何作图、虚拟尺规联动、清屏 / 撤销",
             "Canvas2D、自定义工具组件、动量动画"],
            ["函数探究", "function", "表达式编译、滑动条参数、零点 / 极值标注",
             "自研 Shunting-Yard + 数值方法"],
            ["3D 计算器", "calculator3d", "几何化数学运算演示", "@react-three/drei、Three.js"],
            ["反代后端", "—", "API Key 隐藏、HMAC 鉴权、限流、Prometheus 指标",
             "Rust + axum + tower_governor"],
            ["桌面壳", "—", "本地模型仓库、AI 命令、自绘标题栏、双窗口入口",
             "Tauri 2.x + reqwest"],
        ])

    # ============ 2. 总体架构 ============
    doc.add_page_break()
    add_heading(doc, "2  总体架构", 1)

    add_heading(doc, "2.1  分层视图", 2)
    add_para(doc,
        "系统遵循“前端富客户端 + 反代隔离层 + 第三方 AI 网关 + 桌面壳”四层结构。"
        "Web 与桌面端共享同一份 React 产物（dist/），区别仅在于运行容器与对外通信路径："
        "Web 端经 Nginx → Rust 反代 → 上游 Gemini，桌面端经 Tauri Rust 进程直接访问外网。")

    code = (
        "┌────────────────────────────────────────────────────────────────────┐\n"
        "│                       用户交互层 (Browser / WebView2)              │\n"
        "│                                                                    │\n"
        "│   index.html  门户首页 (原生 JS+CSS, portal.js)                    │\n"
        "│       │ 启动                                                       │\n"
        "│       ▼                                                            │\n"
        "│   app.html    主程序 (React 19 + R3F + Tailwind v4 + Zustand)      │\n"
        "│       ├── AR 模式  ── MediaPipe HandLandmarker (GPU)               │\n"
        "│       ├── 白板    ── Canvas2D + 虚拟尺规组件                       │\n"
        "│       ├── 函数探究 ── 自研表达式引擎                                │\n"
        "│       └── 3D 计算器 ── React-Three-Fiber                            │\n"
        "└──────────┬─────────────────────────────────────────────┬───────────┘\n"
        "           │ Web 部署                                    │ 桌面部署\n"
        "           ▼                                             ▼\n"
        "┌──────────────────────────────┐         ┌──────────────────────────────┐\n"
        "│ Nginx (HTTPS, 静态 + 反代)   │         │ Tauri 2.x WebView2 (Rust)    │\n"
        "│  /assets/* → dist            │         │  ai.rs   ── 直连上游 AI       │\n"
        "│  /api/*    → 127.0.0.1:8787  │         │  local_store.rs ── 本地模型   │\n"
        "└──────────┬───────────────────┘         │  TitleBar 自绘标题栏          │\n"
        "           ▼                              └──────────────────────────────┘\n"
        "┌──────────────────────────────┐\n"
        "│ hologrip-proxy (Rust + axum) │\n"
        "│  HMAC token 鉴权              │\n"
        "│  tower_governor 限流          │\n"
        "│  Prometheus /metrics          │\n"
        "└──────────┬───────────────────┘\n"
        "           ▼\n"
        "┌──────────────────────────────┐\n"
        "│ 上游 AI 网关 (api.gemai.cc)  │\n"
        "│  Gemini-3.5-flash 多模态      │\n"
        "└──────────────────────────────┘"
    )
    add_code(doc, code)

    add_heading(doc, "2.2  关键设计原则", 2)
    add_bullets(doc, [
        "Web / 桌面同源：用 Vite 多入口（portal + app）+ platform.ts 运行时探测，"
        "实现一份代码两种容器，同时让桌面专属的 @tauri-apps/api 调用被 tree-shaking 剔除出 Web bundle。",
        "密钥永不出仓：Web 端通过反代隔离 API Key，桌面端通过 Rust 进程持有 Key，"
        "前端 bundle 任何分支都搜不到 sk-* 真实密钥。",
        "全局状态单一真相源：所有跨组件状态（手势、画笔、模型、悬浮窗）集中在 Zustand store，"
        "组件只声明依赖切片，避免事件总线带来的耦合。",
        "实时性优先：手部追踪、模型旋转、笔迹绘制全部在 requestAnimationFrame 主循环中完成，"
        "并采用预分配对象 + 卡尔曼滤波 + One-Euro 滤波三重平滑，保证 60 FPS 不抖动。",
        "桌面端零安装依赖：Tauri 编译产物为单一 NSIS exe（约 12 MB），不再绑定 Node.js。",
    ])

    # ============ 3. 前端架构 ============
    doc.add_page_break()
    add_heading(doc, "3  前端架构", 1)

    add_heading(doc, "3.1  技术栈", 2)
    add_table(doc,
        ["层次", "组件", "版本", "职责"],
        [
            ["运行时", "React", "19.x", "组件树渲染与并发更新"],
            ["状态管理", "Zustand", "5.x", "全局可订阅状态，浅比较触发渲染"],
            ["渲染引擎", "@react-three/fiber + drei", "9.x / 10.x", "声明式 Three.js / 三维场景"],
            ["3D 内核", "three.js", "0.184", "几何、材质、四元数、射线投射"],
            ["视觉识别", "@mediapipe/tasks-vision", "0.10.34", "HandLandmarker GPU 推理"],
            ["样式", "TailwindCSS", "4.1", "原子化 CSS + 暗色主题"],
            ["动画", "motion", "12.x", "悬浮窗/Dock 微交互"],
            ["构建", "Vite", "6.2", "多入口打包 + HMR"],
            ["类型", "TypeScript", "5.8", "全量静态类型"],
            ["桌面壳", "@tauri-apps/api", "2.x", "桌面端命令通信"],
        ])

    add_heading(doc, "3.2  目录结构", 2)
    add_code(doc,
        "src/\n"
        "├── App.tsx                    应用主壳 + 摄像头 / 手部追踪主循环\n"
        "├── main.tsx                   React 挂载入口\n"
        "├── store.ts                   Zustand 全局状态\n"
        "├── index.css                  Tailwind 入口与暗色变量\n"
        "├── components/\n"
        "│   ├── AppleDock.tsx          底部 Dock 切换栏\n"
        "│   ├── Calculator3D.tsx       三维空间计算器\n"
        "│   ├── Canvas2D.tsx           AR 模式下 2D 笔迹画布\n"
        "│   ├── CameraPermissionModal.tsx 摄像头授权弹窗\n"
        "│   ├── FloatingWindow.tsx     可拖拽悬浮窗容器\n"
        "│   ├── FunctionExplorer.tsx   函数探究面板\n"
        "│   ├── GeometryBoard.tsx      几何作图白板\n"
        "│   ├── MathKeyboard.tsx       数学输入软键盘\n"
        "│   ├── MathModel.tsx          AR 三维模型与表面笔迹\n"
        "│   ├── OverlayUI.tsx          AR 控制条与提示\n"
        "│   ├── ToolboxPanel.tsx       作图工具控制条\n"
        "│   ├── WhiteboardCanvas.tsx   全局穿透书写画布\n"
        "│   ├── desktop/TitleBar.tsx   桌面端自绘标题栏\n"
        "│   └── tools/                 虚拟尺、三角板、量角器、圆规\n"
        "└── lib/\n"
        "    ├── auth.ts                反代 token 生命周期\n"
        "    ├── gemini.ts              AI 几何识图客户端\n"
        "    ├── geometry.ts            顶点归一化 / 三角化\n"
        "    ├── handTracking.ts        手部追踪 + 卡尔曼滤波\n"
        "    ├── mathExpression.ts      Shunting-Yard 表达式引擎\n"
        "    ├── mediapipe.ts           HandLandmarker 工厂\n"
        "    ├── platform.ts            Tauri / Web 运行时探测\n"
        "    ├── rotation.ts            Arcball + One-Euro 旋转算法\n"
        "    └── utils.ts               通用工具函数")

    add_heading(doc, "3.3  全局状态模型", 2)
    add_para(doc,
        "全局 Zustand store 切分为五个语义域：手部状态、模型状态、UI 状态、笔迹状态、主题状态。"
        "每个组件仅订阅自己关心的字段，避免无关变化触发重渲染。下表列出关键字段：")
    add_table(doc,
        ["字段", "类型", "说明"],
        [
            ["leftHand / rightHand", "HandState", "经手部追踪层平滑后的左右手快照"],
            ["activeModel / activeCustomModelId", "MathShape | string | null", "当前展示的内置或自定义模型"],
            ["customModels", "CustomModel[]", "AI 解析得到的用户模型集合"],
            ["surfaceStrokes", "SurfaceStroke[]", "三维模型表面的局部坐标系笔迹"],
            ["modelLines / activeLineStart", "数组 / 端点", "顶点连线模式状态"],
            ["isPenActive / isLineDrawingActive", "boolean", "互斥的写字 / 连线模式开关"],
            ["activeTab", "AppTab", "白板 / 函数 / 计算器 / AR 之间的切换"],
            ["showRuler / showCompass / ...", "boolean", "虚拟作图工具的显隐"],
            ["theme", "'dark' | 'light'", "用户主题；AR 场景始终强制暗色"],
        ])

    # ============ 4. 手部追踪 ============
    doc.add_page_break()
    add_heading(doc, "4  手部追踪与手势交互", 1)

    add_heading(doc, "4.1  数据流水线", 2)
    add_para(doc,
        "AR 模式下，App.tsx 维护一个 requestAnimationFrame 主循环：从 video 取帧 → "
        "MediaPipe HandLandmarker 推理 → 拼装 RawHandObservation → 进入 HandTracker 平滑层 → "
        "写入 Zustand 的 leftHand / rightHand → 触发组件订阅与合成点击。")
    add_code(doc,
        "video frame\n"
        "    │\n"
        "    ▼\n"
        "HandLandmarker.detectForVideo()      ← MediaPipe Tasks Vision (GPU)\n"
        "    │  landmarks (21 关键点 × N 只手) + handedness 投票\n"
        "    ▼\n"
        "RawHandObservation[]                ← App.tsx 计算 NDC / 像素 / 捏合距离\n"
        "    │\n"
        "    ▼\n"
        "HandTracker.update()                ← lib/handTracking.ts\n"
        "    │  匹配 → coast 外推 → 卡尔曼滤波 → 主用户锁定 → 槽位分配\n"
        "    ▼\n"
        "{ left, right }: TrackedHandSnapshot\n"
        "    │\n"
        "    ▼\n"
        "Zustand updateHands  → 各组件订阅 → 合成 click / 旋转 / 缩放")

    add_heading(doc, "4.2  HandTracker 关键算法", 2)
    add_bullets(doc, [
        "常加速度卡尔曼滤波：分别在 X / Y 轴维护位置-速度-加速度 3×3 状态向量与协方差矩阵，"
        "对每帧观测做预测-更新闭环，输出方差约束下的最优估计。",
        "自适应过程噪声 Q：根据当前估计速度自动放大 Q（高速跟手）或缩小 Q（低速去抖）。",
        "Coast 滑行：识别短时丢失（默认 250 ms 内）时使用速度 + 加速度做物理外推，"
        "避免模型瞬间跳跃；外推距离超出阈值时主动放弃，不触发新的合成点击。",
        "主用户锁定：以最先出现的双手作为主用户，旁观者的手会被忽略；"
        "锁定超时 1500 ms 后才允许接管，避免多人镜头下控制权乱跳。",
        "Handedness 投票窗口：环形缓冲 10 帧投票决定 left / right；同时为代价函数加入"
        "handedness 不一致惩罚，杜绝双手交叉时左右手翻转。",
        "槽位锁定：左右手输出与 trackId 绑定，仅在 track 死亡时才解绑，"
        "解决了上一版“双手在屏幕中线时光标乱跳”的问题。",
        "施密特触发器：捏合判定使用 0.045 / 0.06 双阈值，避免边界抖动。",
        "微动死区：估计速度 < 0.12 NDC/s 且位移 < 0.0018 NDC 时强制锁定光标，"
        "彻底消除“手不动光标却在抖”的肌电级噪声。",
    ])

    add_heading(doc, "4.3  Arcball 旋转与缩放", 2)
    add_para(doc,
        "三维模型旋转使用 Bell’s Trackball 投影（Holroyd 改进版）将屏幕 NDC 投到单位球，"
        "再以 a × b 为旋转轴、acos(a·b) 为旋转角生成增量四元数。该算法路径无关，绕一圈回到原点；"
        "并配合 One-Euro Filter 做输入端抖动抑制，低速强滤波吃噪声、高速弱滤波保留跟手感。")
    add_table(doc,
        ["参数", "默认值", "作用"],
        [
            ["arcballGain", "1.0", "Arcball 灵敏度倍率"],
            ["slerpHalfLife", "0.06 s", "四元数球面插值的视觉半衰期"],
            ["maxAngleStepPerFrame", "π/4", "单帧旋转上限，抑制识别跳变"],
            ["cursorDeadzone", "0.0005 NDC", "光标位移死区"],
            ["pinchDistDeadzone", "0.0001 NDC", "双手间距变化死区"],
            ["scaleGain", "4.0", "捏合距离 → 缩放倍率"],
            ["oneEuroMinCutoff", "0.8 Hz", "One-Euro 静止截止频率"],
            ["oneEuroBeta", "0.7", "One-Euro 速度耦合系数"],
        ])

    # ============ 5. AI 几何识图 ============
    doc.add_page_break()
    add_heading(doc, "5  AI 几何识图链路", 1)

    add_heading(doc, "5.1  调用流程", 2)
    add_para(doc,
        "前端 lib/gemini.ts 暴露 parseGeometryImage(base64, mimeType) 函数，"
        "将题目截图与精心设计的系统提示词一起送入 Gemini 多模态模型，"
        "模型按 responseSchema 强约束返回结构化 JSON：")
    add_code(doc,
        "{\n"
        "  reasoning: string,        // 测量与剔除的思考过程\n"
        "  name: string,             // 几何体名称, 如 \"四棱锥 P-ABCD\"\n"
        "  vertices: [ { label, x, y, z } ],\n"
        "  faces: [ [int, int, ...] ],\n"
        "  edges: [ [int, int] ]\n"
        "}")
    add_para(doc,
        "前端通过 lib/geometry.ts 的 normalizeVertices 与 triangulateFaces 把任意比例的"
        "顶点集合归一化到 [-1.5, 1.5] 立方体内，并将多边形面扇形拆分为三角形索引，"
        "供 Three.js BufferGeometry 直接渲染。")

    add_heading(doc, "5.2  双通道适配", 2)
    add_table(doc,
        ["运行模式", "判定条件", "鉴权方式", "适用场景"],
        [
            ["反代模式", "VITE_GEMINI_BASE_URL 以 / 开头", "首次拉取 1h HMAC token，401 自动重签", "Web 部署，bundle 不含 Key"],
            ["直连模式", "VITE_GEMINI_BASE_URL 是 https:// 绝对地址", "Authorization: Bearer ${VITE_GEMINI_API_KEY}", "本地开发"],
            ["桌面命令", "Tauri 容器内", "Rust ai.rs 持有 Key，前端通过 invoke 调用", "桌面安装包"],
        ])

    add_heading(doc, "5.3  提示词设计要点", 2)
    add_bullets(doc, [
        "明确告知模型“图像采用斜二测画法”，给出 X/Y/Z 轴在画面中的对应关系与深度缩放规则。",
        "强制剔除辅助线、动点（如 P、F）与底面对角线，避免污染棱集合。",
        "要求严格还原比例而非套用标准正方体，并给出坐标范围 [-1.5, 1.5]。",
        "使用 responseSchema + responseMimeType 双保险，"
        "让模型只能产出合法 JSON，无 markdown 包装与多余文本。",
    ])

    # ============ 6. 数学引擎 ============
    doc.add_page_break()
    add_heading(doc, "6  数学表达式与函数探究", 1)

    add_heading(doc, "6.1  Shunting-Yard 编译流水线", 2)
    add_para(doc,
        "lib/mathExpression.ts 实现完整的词法分析 → 隐式乘法插入 → Shunting-Yard 转 RPN → "
        "RPN 验证 → 闭包求值的编译流水线，支持中括号与中文 π，"
        "并对一元运算、函数变长参数、右结合 ^ 做了精确处理。")
    add_table(doc,
        ["能力", "示例", "说明"],
        [
            ["四则与幂", "2*x^2 - 3", "完整运算符优先级与右结合"],
            ["隐式乘法", "2x, 3sin(x), (x+1)(x-1)", "解析阶段自动插入 *"],
            ["一元函数", "sin cos tan exp ln log sqrt abs floor ...", "覆盖中学到大学一年级常用函数"],
            ["二元函数", "max(a,b), atan2(y,x), mod(n,m)", "通过 argCount 跟踪元数"],
            ["常量", "pi, π, e, tau", "自动替换为 Math 常量"],
            ["数值方法", "findRoots, findExtrema, numericDerivative", "二分变号 / 中心差分"],
        ])

    add_heading(doc, "6.2  函数探究面板交互", 2)
    add_bullets(doc, [
        "解析失败时降级为 NaN 并保持渲染，避免输入错误时画面空白。",
        "支持滑动条参数注入（如 a、b），曲线随参数实时重绘。",
        "极值与零点采样分辨率默认 600~800 段，二分细化到 1e-12 精度。",
        "前端编译期通过 validateRPN 模拟堆栈，提前发现“缺少参数”等用户错误。",
    ])

    # ============ 7. 反代后端 ============
    doc.add_page_break()
    add_heading(doc, "7  反代后端服务（Rust + axum）", 1)

    add_heading(doc, "7.1  模块划分", 2)
    add_table(doc,
        ["模块", "文件", "职责"],
        [
            ["启动装配", "server/src/main.rs", "环境配置、CORS、限流、路由、优雅退出"],
            ["配置层", "server/src/config.rs", "环境变量 → 强类型 Config 结构"],
            ["鉴权", "server/src/auth.rs", "HMAC-SHA256 token 签发 / 校验 / 后台清扫"],
            ["反代", "server/src/proxy.rs", "鉴权中间件 + 透明转发 handler"],
            ["指标", "server/src/metrics.rs", "Prometheus 指标定义与 /metrics 端点"],
        ])

    add_heading(doc, "7.2  路由与安全策略", 2)
    add_table(doc,
        ["路由", "方法", "鉴权", "说明"],
        [
            ["/healthz", "GET", "公开", "systemd / 监控用健康检查"],
            ["/api/auth/issue", "POST", "Origin 白名单", "签发 1h 短期 token，绑定客户端 IP，限调用次数"],
            ["/api/gemini/{*path}", "ANY", "Bearer token", "透明转发到 UPSTREAM_BASE_URL，注入真实 Key"],
            ["/metrics", "GET", "仅监听 127.0.0.1:9898", "Prometheus 抓取，不暴露公网"],
        ])

    add_heading(doc, "7.3  关键中间件链", 2)
    add_code(doc,
        "Router\n"
        "  └ /healthz\n"
        "  └ /api/auth/issue\n"
        "  └ /api/gemini/* (require_token middleware)\n"
        "  ── GovernorLayer            (per-IP 令牌桶限流)\n"
        "  ── RequestBodyLimitLayer    (默认 16 MB)\n"
        "  ── DefaultBodyLimit::max    (与上层对齐)\n"
        "  ── CorsLayer                (Origin 白名单 + 放行 Authorization)\n"
        "  └ TraceLayer                (结构化日志)")

    add_heading(doc, "7.4  Prometheus 指标", 2)
    add_table(doc,
        ["指标", "类型", "标签", "用途"],
        [
            ["proxy_requests_total", "Counter", "path, status", "QPS / 状态码分布"],
            ["proxy_upstream_duration_seconds", "Histogram", "path", "上游耗时直方图"],
            ["proxy_in_flight_requests", "Gauge", "—", "在飞请求数"],
            ["proxy_rate_limited_total", "Counter", "—", "被限流的请求数"],
            ["auth_token_issued_total", "Counter", "—", "签发的 token 数"],
            ["auth_token_rejected_total", "Counter", "reason", "鉴权失败原因细分"],
        ])

    # ============ 8. 桌面端 ============
    doc.add_page_break()
    add_heading(doc, "8  桌面端架构（Tauri 2.x）", 1)

    add_heading(doc, "8.1  窗口与入口", 2)
    add_para(doc,
        "桌面端启动时加载 app.html 作为唯一窗口，标题栏关闭（decorations: false），"
        "由 React 端的 components/desktop/TitleBar.tsx 自绘。"
        "窗口尺寸 1280×820，最小 960×640，背景透明并启用 OS 阴影。"
        "命令 open_simulation_window 支持以 label = simulation 复用同一窗口。")

    add_heading(doc, "8.2  Tauri 命令清单", 2)
    add_table(doc,
        ["命令", "参数", "返回", "说明"],
        [
            ["parse_geometry_image", "image_base64, mime_type", "JSON", "调上游 AI 解析几何"],
            ["list_ai_models", "—", "Vec<JSON>", "列出本地保存的模型"],
            ["save_ai_model", "{ model: JSON }", "{ file_name, path }", "持久化到 app_data_dir/ai-models"],
            ["delete_ai_model", "{ id }", "()", "按 id 删除单个模型"],
            ["clear_ai_models", "—", "()", "清空本地模型仓库"],
            ["ai_models_dir", "—", "String", "返回模型仓库路径"],
            ["open_simulation_window", "—", "()", "复用 / 新建仿真窗口"],
        ])

    add_heading(doc, "8.3  本地模型仓库", 2)
    add_para(doc,
        "src-tauri/src/local_store.rs 把 AI 解析结果以 JSON 文件形式存放到 "
        "app_data_dir/ai-models/。文件名含时间戳与短哈希以便自然排序，"
        "重启时由前端逐个回灌到 Zustand 的 customModels。"
        "由于 Tauri 沙箱限制，PrivateTmp 与 ProtectSystem 环境下读写仅限该目录。")

    add_heading(doc, "8.4  打包与分发", 2)
    add_bullets(doc, [
        "构建链：npm run build → vite 多入口产物 → tauri build → NSIS 打包。",
        "默认安装模式 perMachine，安装包语言锁定简体中文，自定义图标与启动器横幅。",
        "WebView2 启用 --use-fake-ui-for-media-stream，"
        "使桌面端摄像头授权静默通过，避免反复弹出系统询问。",
        "桌面端通过 lib/platform.ts 的 isDesktop 在运行时分流，"
        "tree-shaking 后 Web bundle 不会包含任何 Tauri API。",
    ])

    # ============ 9. 部署与运维 ============
    doc.add_page_break()
    add_heading(doc, "9  部署与运维", 1)

    add_heading(doc, "9.1  Web 端部署拓扑", 2)
    add_code(doc,
        "用户浏览器\n"
        "    │\n"
        "    ▼ HTTPS\n"
        "Nginx (宝塔面板)\n"
        "    ├── /          → /www/wwwroot/<域名>/index.html\n"
        "    ├── /app.html  → /www/wwwroot/<域名>/app.html\n"
        "    ├── /assets/*  → 静态缓存 30 天\n"
        "    └── /api/*     → http://127.0.0.1:8787  (proxy_buffering off)\n"
        "                                    │\n"
        "                                    ▼\n"
        "                         systemd: hologrip-proxy.service\n"
        "                                    │\n"
        "                                    ▼\n"
        "                         上游 AI 网关 (api.gemai.cc)")

    add_heading(doc, "9.2  关键 systemd 配置", 2)
    add_code(doc,
        "[Service]\n"
        "Type=simple\n"
        "WorkingDirectory=/opt/hologrip-proxy\n"
        "EnvironmentFile=/opt/hologrip-proxy/.env\n"
        "ExecStart=/opt/hologrip-proxy/target/release/hologrip-proxy\n"
        "Restart=on-failure\n"
        "NoNewPrivileges=true\n"
        "PrivateTmp=true\n"
        "ProtectSystem=strict\n"
        "ProtectHome=true\n"
        "ReadWritePaths=/opt/hologrip-proxy")

    add_heading(doc, "9.3  关键环境变量", 2)
    add_table(doc,
        ["变量", "示例", "必填", "用途"],
        [
            ["UPSTREAM_BASE_URL", "https://api.gemai.cc", "✓", "上游 AI 网关"],
            ["UPSTREAM_API_KEY", "sk-***", "✓", "真实 Key，仅服务端持有"],
            ["AUTH_HMAC_SECRET", "openssl rand -hex 32", "✓", "token 签名密钥（≥32 字节）"],
            ["AUTH_TOKEN_TTL_SECS", "3600", "—", "token 有效期"],
            ["AUTH_TOKEN_QUOTA", "100", "—", "单 token 调用次数上限"],
            ["RATE_LIMIT_PER_SECOND", "2", "—", "每 IP 每秒补充令牌数"],
            ["RATE_LIMIT_BURST", "10", "—", "每 IP 突发桶容量"],
            ["MAX_BODY_BYTES", "16777216", "—", "请求体上限"],
            ["CORS_ALLOWED_ORIGINS", "https://your-domain", "—", "浏览器 CORS 白名单"],
            ["AUTH_ISSUE_ALLOWED_ORIGINS", "https://your-domain", "—", "签发 token Origin 白名单"],
        ])

    add_heading(doc, "9.4  灾难恢复与可观测性", 2)
    add_bullets(doc, [
        "服务重启即吊销所有旧 token：HMAC 私钥常驻进程内存，前端 401 后自动重签发。",
        "tower_governor 限流采用 SmartIpKeyExtractor，反代场景下优先读 X-Forwarded-For。",
        "/metrics 端点仅在 127.0.0.1:9898 监听，不通过 Nginx 暴露公网，建议通过 SSH 隧道抓取。",
        "前端在反代模式下捕获 401 自动作废本地 token 并重试一次，对用户完全无感。",
        "桌面端 ai_models_dir 命令可一键定位本地模型仓库，便于排障与备份。",
    ])

    # ============ 10. 安全 ============
    doc.add_page_break()
    add_heading(doc, "10  安全与合规设计", 1)

    add_bullets(doc, [
        "API Key 隔离：Web 端通过反代隔离 Gemini Key，桌面端通过 Rust 进程隔离，"
        "前端 bundle 在任何分支都不持有真实 Key。",
        "鉴权链：HMAC-SHA256 token + IP 绑定 + quota 配额 + TTL 过期，"
        "服务重启自动作废，副本水平扩展时通过共享 HMAC 密钥保证一致性。",
        "请求体限制：DefaultBodyLimit + RequestBodyLimitLayer 双重 16 MB 上限，"
        "防御大文件灌入与内存膨胀。",
        "CORS 白名单：默认允许 * 仅供本地开发；生产配置必须显式列出 https 域名，"
        "并放行 Authorization 头让浏览器 preflight 通过。",
        "摄像头授权：通过 CameraPermissionModal 显式向用户解释用途，"
        "并将授权结果写入 localStorage，避免反复弹窗。",
        "桌面端沙箱：systemd 中启用 NoNewPrivileges、ProtectSystem=strict、ProtectHome、PrivateTmp，"
        "ReadWritePaths 限制仅写入工作目录。",
        "依赖审计：Rust 端启用 codegen-units=1 + LTO + strip 减小符号表，"
        "前端用 Vite 静态产物便于 SAST 扫描。",
    ])

    # ============ 11. 性能 ============
    doc.add_page_break()
    add_heading(doc, "11  性能与体验优化", 1)

    add_table(doc,
        ["优化点", "措施", "收益"],
        [
            ["手部追踪稳定性", "卡尔曼 + One-Euro + 死区 + Coast 外推", "彻底消除静止抖动与短时丢失跳变"],
            ["旋转无路径依赖", "Holroyd 球面投影 Arcball", "绕一圈回到原点，无累计漂移"],
            ["合成点击精度", "AR 舞台 ref + 视口偏移补偿", "桌面端 36px 自绘标题栏不再造成顶部按钮点不到"],
            ["三维渲染", "Environment + Lightformer 烘焙反射", "无需运行时光源采样仍获得 PBR 效果"],
            ["状态触发", "Zustand 浅比较 + 引用替换", "AR 模式下 60FPS 不丢帧"],
            ["反代延迟", "reqwest stream + proxy_buffering off", "Gemini 流式响应零首字延迟"],
            ["桌面包体", "Cargo strip + LTO + NSIS 压缩", "单文件安装包约 12 MB"],
            ["前端首屏", "Vite 多入口 + 静态资源 immutable 缓存", "门户秒开，应用按需加载"],
        ])

    # ============ 12. 后续演进 ============
    doc.add_page_break()
    add_heading(doc, "12  后续演进规划", 1)

    add_bullets(doc, [
        "多人协作：在反代层引入 WebSocket 通道，让同一房间内的师生共享白板与模型操作。",
        "向量化笔迹：将 Canvas 像素笔迹升级为 SVG / GeoJSON，支持回放、搜索与导出 PDF。",
        "AI 解题助教：在 AI 几何识图基础上扩展“证明步骤生成”，输出可点开的推导树。",
        "VR 头显：通过 WebXR 把当前 AR 模式无缝迁移到 Meta Quest / Vision Pro。",
        "数据分析后台：在 Prometheus 之上接入 Grafana + Loki，"
        "为运营提供活跃用户与失败链路的全景看板。",
        "国产替代：将 MediaPipe 替换为本地化 ONNX 模型，使桌面端在完全离线场景下仍可工作。",
    ])

    add_para(doc,
        "本文档随代码一同维护，最新版本以仓库主分支为准。",
        size_pt=10, indent_first=False)

    doc.save(OUT)
    print(f"已生成: {OUT}")


if __name__ == "__main__":
    build()
