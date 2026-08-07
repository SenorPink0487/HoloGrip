import os
import docx
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml import OxmlElement, parse_xml
from docx.oxml.ns import qn, nsdecls

doc = docx.Document()

# 页边距设置
for section in doc.sections:
    section.top_margin = Inches(0.8)
    section.bottom_margin = Inches(0.8)
    section.left_margin = Inches(0.8)
    section.right_margin = Inches(0.8)

# 默认字体设置
style = doc.styles['Normal']
style.font.name = 'Microsoft YaHei'
style.font.size = Pt(10.5)
style._element.rPr.rFonts.set(qn('w:eastAsia'), 'Microsoft YaHei')

# 大标题
title_p = doc.add_paragraph()
title_p.paragraph_format.space_before = Pt(0)
title_p.paragraph_format.space_after = Pt(14)
title_run = title_p.add_run('现有教学 / 教辅类产品功能分析表（简明修正版）')
title_run.font.name = 'Microsoft YaHei'
title_run.font.size = Pt(18)
title_run.font.bold = True
title_run.font.color.rgb = RGBColor(0x2B, 0x3A, 0x4A)
title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER

# 创建表格 (7 行 3 列)
table = doc.add_table(rows=7, cols=3)
table.alignment = WD_TABLE_ALIGNMENT.CENTER
table.autofit = False

col_widths = [Inches(1.6), Inches(2.7), Inches(2.7)]
headers = ['现有产品', '它能做什么', '它做不到什么']

def set_cell_background(cell, fill_hex):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{fill_hex}"/>')
    tcPr.append(shd)

def set_cell_margins(cell, top=100, bottom=100, left=140, right=140):
    tcPr = cell._tc.get_or_add_tcPr()
    tcMar = OxmlElement('w:tcMar')
    for m, val in [('top', top), ('bottom', bottom), ('left', left), ('right', right)]:
        node = OxmlElement(f'w:{m}')
        node.set(qn('w:w'), str(val))
        node.set(qn('w:type'), 'dxa')
        tcMar.append(node)
    tcPr.append(tcMar)

def set_table_borders(table, color="D1D5DB", sz="6", val="single"):
    tblPr = table._tbl.tblPr
    borders = parse_xml(
        f'<w:tblBorders {nsdecls("w")}>\n'
        f'  <w:top w:val="{val}" w:sz="{sz}" w:space="0" w:color="{color}"/>\n'
        f'  <w:left w:val="none"/>\n'
        f'  <w:bottom w:val="{val}" w:sz="{sz}" w:space="0" w:color="{color}"/>\n'
        f'  <w:right w:val="none"/>\n'
        f'  <w:insideH w:val="{val}" w:sz="{sz}" w:space="0" w:color="{color}"/>\n'
        f'  <w:insideV w:val="{val}" w:sz="{sz}" w:space="0" w:color="{color}"/>\n'
        f'</w:tblBorders>'
    )
    tblPr.append(borders)

set_table_borders(table, color="D1D5DB", sz="6")

# 格式化表头 (深蓝灰背景 + 白色粗体)
hdr_cells = table.rows[0].cells
for i, header_text in enumerate(headers):
    hdr_cells[i].width = col_widths[i]
    set_cell_background(hdr_cells[i], '2B3A4A')
    set_cell_margins(hdr_cells[i], top=140, bottom=140, left=140, right=140)
    hdr_cells[i].vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    
    p = hdr_cells[i].paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(3)
    p.paragraph_format.space_after = Pt(3)
    
    run = p.add_run(header_text)
    run.font.name = 'Microsoft YaHei'
    run.font.size = Pt(11)
    run.font.bold = True
    run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)

header_tr = table.rows[0]._tr.get_or_add_trPr()
header_tr.append(parse_xml(f'<w:tblHeader {nsdecls("w")}/>'))

# 精简概括版数据 (修复文字笔误)
data = [
    (
        "GeoGebra 3D\n(动态数学)",
        "• 3D 隐/显式曲面与空间几何解算\n• 庞大开源社区与多端支持\n• 移动端基础 AR 投射",
        "• 无物理/化学等跨学科仿真引擎\n• 依赖鼠标触屏，无摄像头裸眼手势\n• 无 AI 2D 题目 3D 拓扑重构"
    ),
    (
        "NOBOOK (NB实验室)\n(K12 虚拟仿真)",
        "• 同步 K12 课本物理、化学经典实验\n• 免器材耗材与安全隐患，自由组装\n• 支持 PC/移动/电子白板多端演示",
        "• 局限于 2D/2.5D 平面交互，缺乏 3D 空间计算\n• 标准拖拽交互，无手势捏合/旋转抓取\n• 缺乏高阶科研级物理（航天、流体）探究"
    ),
    (
        "PhET\n(科学互动仿真)",
        "• 基础数理化概念探究与可视化演示\n• 微观/宏观粒子互动调节\n• 纯 Web 轻量化运行，免费开源",
        "• 界面多为二维示意，无 3D 建模渲染\n• 无法进行复杂公式推导与深度方程解算\n• 无 AI 解题与课件生成辅助"
    ),
    (
        "Merge Cube & JigSpace\n(AR 3D 概念展示)",
        "• Merge：实体黑方块“手握”3D 模型\n• JigSpace：工业与科学结构步骤拆解演示\n• 支持高级 AR 设备交互",
        "• 依赖实体方块教具或昂贵头显(Vision Pro)\n• 偏向静态模型展示，无公式物理解算\n• 缺乏沙盒自由创作环境"
    ),
    (
        "希沃 (Seewo) 白板\n(K12 交互白板生态)",
        "• 国内 K12 占有率第一，大屏互动极强\n• AI 备课、板书美化与二维题目识别\n• 丰富教研资源与授课工具",
        "• 聚焦 2D 平面白板，无 3D 空间渲染引擎\n• 仅限大屏触控，无三维手势空间追踪\n• 无物理/化学的高精度计算仿真引擎"
    ),
    (
        "Labster\n(高阶虚拟实验室)",
        "• 高保真 3D 虚拟实验室与剧情化实验\n• 涵盖高校生物、化学、医学高阶实验\n• 规范的教学指导与答题验证",
        "• 授权部署成本高，主要面向高校\n• 流程相对固定（按剧情点击），缺乏自由沙盒\n• 无法通过自定义公式驱动场景"
    )
]

for row_idx, (col1, col2, col3) in enumerate(data, start=1):
    row_cells = table.rows[row_idx].cells
    bg_color = 'F9FAFB' if row_idx % 2 == 1 else 'FFFFFF'
    
    trPr = table.rows[row_idx]._tr.get_or_add_trPr()
    trPr.append(parse_xml(f'<w:cantSplit {nsdecls("w")}/>'))
    
    for c_idx, content in enumerate([col1, col2, col3]):
        cell = row_cells[c_idx]
        cell.width = col_widths[c_idx]
        set_cell_background(cell, bg_color)
        set_cell_margins(cell, top=100, bottom=100, left=140, right=140)
        cell.vertical_alignment = WD_ALIGN_VERTICAL.TOP
        
        p = cell.paragraphs[0]
        p.paragraph_format.space_before = Pt(2)
        p.paragraph_format.space_after = Pt(2)
        p.paragraph_format.line_spacing = 1.15
        
        if c_idx == 0:
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            lines = content.split('\n')
            for l_idx, line in enumerate(lines):
                if l_idx > 0:
                    p = cell.add_paragraph()
                    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    p.paragraph_format.space_before = Pt(2)
                    p.paragraph_format.space_after = Pt(2)
                    p.paragraph_format.line_spacing = 1.15
                
                run = p.add_run(line)
                run.font.name = 'Microsoft YaHei'
                run.font.size = Pt(10)
                run.font.bold = True
                run.font.color.rgb = RGBColor(0x1F, 0x29, 0x37)
        else:
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            lines = content.split('\n')
            for l_idx, line in enumerate(lines):
                if l_idx > 0:
                    p = cell.add_paragraph()
                    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
                    p.paragraph_format.space_before = Pt(2)
                    p.paragraph_format.space_after = Pt(2)
                    p.paragraph_format.line_spacing = 1.15
                
                run = p.add_run(line)
                run.font.name = 'Microsoft YaHei'
                run.font.size = Pt(9.5)
                run.font.color.rgb = RGBColor(0x37, 0x41, 0x51)

output_filename = "d:/math/现有产品分析表_修正版.docx"
doc.save(output_filename)
print(f"Successfully saved to {output_filename}")
